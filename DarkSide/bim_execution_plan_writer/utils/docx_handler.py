import mammoth
from docx import Document

def convert_to_html(filepath):
    try:
        with open(filepath, "rb") as docx_file:
            result = mammoth.convert_to_html(docx_file)
            return result.value
    except Exception as e:
        return f"<p>Error converting file: {str(e)}</p>"

def read_text(filepath):
    try:
        doc = Document(filepath)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        # Also read tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        return '\n'.join(full_text)
    except Exception as e:
        return ""

def save_docx(filepath, edits):
    doc = Document(filepath)
    
    for edit in edits:
        if edit.get('type') == 'replace':
            original = edit.get('original', '')
            new_text = edit.get('new', '')
            if not original: 
                continue
                
            # Try to replace in paragraphs preserving formatting
            replaced = False
            for paragraph in doc.paragraphs:
                if original in paragraph.text:
                    # improved replacement that tries to preserve runs
                    if replace_text_in_paragraph(paragraph, original, new_text):
                        replaced = True
            
            # Try to replace in tables
            if not replaced: # Continue searching even if found in paragraph? Usually yes, but this logic was existing. 
                             # Let's allow multiple replacements per doc if needed.
                pass # The previous logic stopped checking tables if para found. Let's fix that too.
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if original in cell.text:
                            # Recursive call for paragraphs in cell (cells have paragraphs)
                            for paragraph in cell.paragraphs:
                                if original in paragraph.text:
                                    if replace_text_in_paragraph(paragraph, original, new_text):
                                        replaced = True
                                
        elif edit.get('type') == 'append':
            text = edit.get('text', '')
            if text:
                doc.add_paragraph(text)
            
    doc.save(filepath)

def replace_text_in_paragraph(paragraph, original, new_text):
    """
    Attempts to replace text while preserving formatting.
    Strategy:
    1. Check if the exact text exists in a single run. If so, replace it there.
    2. If not, fall back to replacing the whole paragraph text (loses formatting).
    """
    # 1. Try single run replacement (best for formatting)
    for run in paragraph.runs:
        if original in run.text:
            run.text = run.text.replace(original, new_text)
            return True
            
    # 2. Fallback: The text might be split across runs or simply complex.
    # Replacing paragraph.text wipes formatting.
    # To mitigate, we only do this if we really have to.
    if original in paragraph.text:
        paragraph.text = paragraph.text.replace(original, new_text)
        return True
        
    return False
